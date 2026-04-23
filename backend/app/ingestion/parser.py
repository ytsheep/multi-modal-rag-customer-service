from dataclasses import dataclass
from pathlib import Path
import re

import fitz
from docx import Document as DocxDocument

from app.core.config import get_settings
from app.ingestion.ocr import AliyunOCRClient


@dataclass
class ParsedBlock:
    text: str
    page: int | None = None
    block_type: str = "paragraph"
    title_path: str = ""


class DocumentParser:
    def parse(self, path: Path) -> list[ParsedBlock]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix in {".md", ".markdown"}:
            return self._parse_markdown(path)
        if suffix == ".txt":
            return self._parse_text(path)
        raise ValueError(f"暂不支持的文件类型：{suffix}")

    def _parse_pdf(self, path: Path) -> list[ParsedBlock]:
        doc = fitz.open(path)
        toc_by_page = self._toc_by_page(doc)
        page_blocks: list[list[str]] = []
        text_length = 0

        for page in doc:
            blocks = self._page_text_blocks(page)
            page_blocks.append(blocks)
            text_length += sum(len(block.strip()) for block in blocks)

        if len(doc) > 0 and text_length / len(doc) < 30:
            return self._parse_scanned_pdf(path)

        parsed: list[ParsedBlock] = []
        for page_index, blocks in enumerate(page_blocks, start=1):
            title_path = toc_by_page.get(page_index, "")
            for text in blocks:
                for part in self._split_text_block(text):
                    parsed.append(
                        ParsedBlock(
                            text=part,
                            page=page_index,
                            title_path=title_path,
                        )
                    )
        return parsed

    def _parse_scanned_pdf(self, path: Path) -> list[ParsedBlock]:
        settings = get_settings()
        if not settings.aliyun_ocr_enabled:
            raise RuntimeError("检测到扫描件 PDF。请配置 ALIYUN_OCR_ENABLED=true 后使用阿里云 OCR。")

        ocr = AliyunOCRClient()
        blocks: list[ParsedBlock] = []
        doc = fitz.open(path)
        toc_by_page = self._toc_by_page(doc)
        for page_index, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            text = ocr.recognize_png(pix.tobytes("png"))
            for part in self._split_text_block(text):
                blocks.append(
                    ParsedBlock(
                        text=part,
                        page=page_index,
                        title_path=toc_by_page.get(page_index, ""),
                    )
                )
        return blocks

    def _parse_docx(self, path: Path) -> list[ParsedBlock]:
        doc = DocxDocument(path)
        blocks: list[ParsedBlock] = []
        title_stack: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            heading_level = self._docx_heading_level(paragraph)
            if heading_level:
                title_stack = title_stack[: heading_level - 1]
                title_stack.append(text)
                blocks.append(
                    ParsedBlock(
                        text=text,
                        block_type="title",
                        title_path=" > ".join(title_stack),
                    )
                )
            else:
                blocks.append(ParsedBlock(text=text, title_path=" > ".join(title_stack)))

        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append(
                ParsedBlock(
                    text=self._table_to_markdown(rows),
                    block_type="table",
                    title_path=" > ".join(title_stack),
                )
            )
        return blocks

    def _parse_markdown(self, path: Path) -> list[ParsedBlock]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return self._parse_plain_text(text, markdown=True)

    def _parse_text(self, path: Path) -> list[ParsedBlock]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return self._parse_plain_text(text, markdown=False)

    def _parse_plain_text(self, text: str, markdown: bool) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        title_stack: list[str] = []
        parts = re.split(r"\n{2,}", text)
        for part in parts:
            stripped = part.strip()
            if not stripped:
                continue
            level = self._markdown_heading_level(stripped) if markdown else None
            if level:
                title = stripped.lstrip("#").strip()
                title_stack = title_stack[: level - 1]
                title_stack.append(title)
                blocks.append(
                    ParsedBlock(
                        text=title,
                        block_type="title",
                        title_path=" > ".join(title_stack),
                    )
                )
            else:
                blocks.append(ParsedBlock(text=stripped, title_path=" > ".join(title_stack)))
        return blocks

    def _page_text_blocks(self, page: fitz.Page) -> list[str]:
        raw_blocks = page.get_text("blocks", sort=True)
        blocks = []
        for raw in raw_blocks:
            if len(raw) >= 7 and raw[6] != 0:
                continue
            text = str(raw[4] if len(raw) > 4 else "").strip()
            if text:
                blocks.append(text)
        if blocks:
            return blocks
        text = page.get_text("text").strip()
        return [text] if text else []

    def _split_text_block(self, text: str) -> list[str]:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        parts = [part.strip() for part in re.split(r"\n{2,}", normalized) if part.strip()]
        return parts or ([normalized] if normalized else [])

    def _toc_by_page(self, doc: fitz.Document) -> dict[int, str]:
        toc = doc.get_toc(simple=True)
        if not toc:
            return {}

        entries: list[tuple[int, str]] = []
        stack: list[str] = []
        for level, title, page in toc:
            if page <= 0:
                continue
            clean_title = self._clean_title(title)
            if not clean_title:
                continue
            stack = stack[: level - 1]
            stack.append(clean_title)
            entries.append((int(page), " > ".join(stack)))

        page_titles: dict[int, str] = {}
        current = ""
        cursor = 0
        entries.sort(key=lambda item: item[0])
        for page_number in range(1, len(doc) + 1):
            while cursor < len(entries) and entries[cursor][0] <= page_number:
                current = entries[cursor][1]
                cursor += 1
            page_titles[page_number] = current
        return page_titles

    def _clean_title(self, title: str) -> str:
        title = re.sub(r"\s+", " ", title).strip()
        title = re.sub(r"\.{4,}\s*\d*$", "", title).strip()
        return title

    def _docx_heading_level(self, paragraph: object) -> int | None:
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
        if not match:
            return None
        return min(int(match.group(1)), 6)

    def _markdown_heading_level(self, text: str) -> int | None:
        match = re.match(r"^(#{1,6})\s+.+", text)
        return len(match.group(1)) if match else None

    def _table_to_markdown(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = "| " + " | ".join(normalized[0]) + " |"
        sep = "| " + " | ".join(["---"] * width) + " |"
        body = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
        return "\n".join([header, sep, *body])
